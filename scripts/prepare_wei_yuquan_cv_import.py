from __future__ import annotations

import json
import re
from collections import OrderedDict
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs" / "portal-import" / "wei-yuquan-2026-7"
SOURCE_CANDIDATES = [
    Path(r"C:/Users/baoju/Desktop/导入数据/魏雨泉简历 (2026.7).docx"),
    Path(r"C:/Users/baoju/Desktop/魏雨泉简历 (2026.7).docx"),
]

HEADER_FILL = PatternFill("solid", fgColor="EAF5EE")
HEADER_FONT = Font(bold=True, color="1F3D2B")
TITLE_FILL = PatternFill("solid", fgColor="00873C")
TITLE_FONT = Font(bold=True, color="FFFFFF")
THIN_SIDE = Side(style="thin", color="DDE5DF")
CELL_BORDER = Border(left=THIN_SIDE, right=THIN_SIDE, top=THIN_SIDE, bottom=THIN_SIDE)

SECTION_TITLES = [
    "基本情况",
    "工作经历",
    "教育经历",
    "社会兼职",
    "人才类项目",
    "主持项目",
    "参与项目",
    "科研获奖",
    "教学课程",
    "教改项目",
    "教改论文",
    "教学获奖",
    "出版书籍/专著",
    "出版教材",
    "申请专利",
    "代表作",
    "第一作者或通讯作者发表文章 （非农大第一单位标红）",
    "参与发表文章",
    "会议论文",
    "成果鉴定",
    "国际学术会议作大会报告、特邀报告",
    "创业类项目获奖",
    "指导学生获奖",
    "团队展示",
    "指导研究生",
]

PUBLICATION_HEADERS = ["GB/T 7714-2025格式引文", "摘要", "可见范围", "PDF附件", "首页排序"]
PROJECT_HEADERS = ["项目名称", "项目编号", "资助来源", "负责人", "状态", "开始日期", "结束日期", "经费", "可见范围", "说明", "首页排序"]
PATENT_HEADERS = ["专利名称", "专利号", "状态", "发明人", "PDF附件", "申请日期", "授权日期", "可见范围", "首页排序"]
AWARD_HEADERS = ["奖项名称", "奖项等级", "获奖日期", "参与人员", "获奖图片", "附件PDF", "可见范围", "说明", "首页排序"]
MEMBER_HEADERS = ["姓名", "身份头衔", "邮箱", "研究方向", "头像", "简介", "展示排序"]
STUDENT_HEADERS = ["姓名", "学位类型", "年级", "培养情况", "毕业时间", "研究方向/论文题目", "建议导师", "备注"]


def find_source() -> Path:
    for path in SOURCE_CANDIDATES:
        if path.exists():
            return path
    raise FileNotFoundError("未找到《魏雨泉简历 (2026.7).docx》，请确认文件在桌面或桌面/导入数据目录。")


def load_docx(path: Path) -> tuple[list[str], list[list[list[str]]]]:
    document = Document(str(path))
    paragraphs = [" ".join(paragraph.text.split()) for paragraph in document.paragraphs]
    paragraphs = [line for line in paragraphs if line]
    tables: list[list[list[str]]] = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([" ".join(cell.text.split()) for cell in row.cells])
        tables.append(rows)
    return paragraphs, tables


def get_section(paragraphs: list[str], title: str) -> list[str]:
    try:
        start = paragraphs.index(title) + 1
    except ValueError:
        return []
    end = len(paragraphs)
    for marker in SECTION_TITLES:
        if marker == title:
            continue
        try:
            index = paragraphs.index(marker, start)
        except ValueError:
            continue
        end = min(end, index)
    return paragraphs[start:end]


def normalize_date(value: str) -> str:
    text = (value or "").strip()
    text = text.replace("年", "-").replace("月", "-").replace("日", "")
    text = text.replace(".", "-").replace("/", "-").replace("—", "-").replace("–", "-")
    text = re.sub(r"-+", "-", text).strip("-")
    if not text or text in {"至今", "现在"}:
        return ""
    if re.fullmatch(r"\d{4}", text):
        return f"{text}-01-01"
    if re.fullmatch(r"\d{4}-\d{1,2}", text):
        year, month = text.split("-")
        return f"{year}-{int(month):02d}-01"
    if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}", text):
        year, month, day = text.split("-")
        return f"{year}-{int(month):02d}-{int(day):02d}"
    return ""


def parse_date_range(line: str) -> tuple[str, str]:
    prefix = re.split(r"，|,", line, maxsplit=1)[0]
    tokens = re.findall(r"\d{4}(?:[./]\d{1,2})?(?:[./]\d{1,2})?|至今|现在", prefix)
    if not tokens:
        return "", ""
    start = normalize_date(tokens[0])
    end = normalize_date(tokens[1]) if len(tokens) > 1 else ""
    return start, end


def year_month_date(line: str) -> str:
    match = re.search(r"(19\d{2}|20\d{2})年\s*(\d{1,2})月?", line)
    if match:
        return f"{match.group(1)}-{int(match.group(2)):02d}-01"
    match = re.search(r"(19\d{2}|20\d{2})", line)
    return f"{match.group(1)}-01-01" if match else ""


def compact_text(lines: Iterable[str], bullet: bool = True) -> str:
    cleaned = [line.strip().rstrip("；;") for line in lines if line.strip()]
    if bullet:
        return "\n".join(f"- {line}" for line in cleaned)
    return "\n".join(cleaned)


def build_member_profile(paragraphs: list[str]) -> str:
    intro = get_section(paragraphs, "基本情况")
    profile_sections: list[tuple[str, list[str]]] = [
        ("基本情况", intro[:1]),
        ("工作经历", get_section(paragraphs, "工作经历")),
        ("教育经历", get_section(paragraphs, "教育经历")),
        ("学术与社会服务", get_section(paragraphs, "社会兼职")),
        ("人才与团队项目", get_section(paragraphs, "人才类项目")),
        ("教学工作", get_section(paragraphs, "教学课程")),
        ("教改项目与论文", get_section(paragraphs, "教改项目") + get_section(paragraphs, "教改论文")),
        ("出版书籍与教材", get_section(paragraphs, "出版书籍/专著") + get_section(paragraphs, "出版教材")),
        ("国际会议报告", get_section(paragraphs, "国际学术会议作大会报告、特邀报告")),
        ("团队与研究生指导", get_section(paragraphs, "指导研究生")[:1]),
    ]
    blocks = []
    for title, lines in profile_sections:
        if not lines:
            continue
        blocks.append(f"## {title}\n{compact_text(lines)}")
    return "\n\n".join(blocks)


def member_rows(paragraphs: list[str]) -> list[list]:
    intro = get_section(paragraphs, "基本情况")
    summary = intro[0] if intro else "中国农业大学资源与环境学院生态系副教授，博士生导师。"
    return [[
        "魏雨泉",
        "副教授、博士生导师",
        "weiyq2019@cau.edu.cn",
        "微生物生态；有机废弃物资源化；堆肥微生态调控",
        "导入后在成员编辑页上传头像",
        build_member_profile(paragraphs) or summary,
        1,
    ]]


def strip_leading_date(line: str) -> str:
    return re.sub(r"^\s*\d{4}(?:[./]\d{1,2})?(?:[./]\d{1,2})?\s*[-至到—–~]+\s*(?:至今|现在|\d{4}(?:[./]\d{1,2})?(?:[./]\d{1,2})?)，?", "", line).strip()


def extract_project_number(line: str) -> str:
    candidates = re.findall(r"[A-Z]{1,6}[-]?\d{4,}[-A-Z0-9]*|20\d{2}[A-Z]{2,}\d+|NK\d{6,}|\d{8,}[A-Z]?", line)
    if not candidates:
        return ""
    # Prefer explicit grant numbers inside the final parentheses over money/phone-like numbers.
    for candidate in candidates[::-1]:
        if not re.fullmatch(r"20\d{6}", candidate):
            return candidate
    return candidates[-1]


def extract_amount(line: str) -> str:
    match = re.search(r"(\d+(?:\.\d+)?)\s*万", line)
    return match.group(1) if match else ""


def project_role(line: str, hosted: bool) -> str:
    for role in ["课题负责人", "项目负责人", "项目骨干", "参加人", "参与人"]:
        if role in line:
            return role
    return "项目负责人" if hosted else "参与人"


def project_status(end_date: str) -> str:
    if not end_date:
        return "在研"
    return "在研" if end_date >= "2026-01-01" else "已结题"


def project_title_and_source(line: str, hosted: bool) -> tuple[str, str]:
    text = strip_leading_date(line)
    quoted = re.findall(r"“([^”]+)”", text)
    source = ""
    title = ""
    if quoted:
        title = quoted[0]
        before = text.split("“", 1)[0].strip("，, ")
        parts = [part.strip() for part in re.split(r"，|,", before) if part.strip()]
        source = parts[-1] if parts else ""
        if not source:
            after_quote = text.split("”", 1)[1] if "”" in text else ""
            after_parts = [part.strip() for part in re.split(r"，|,", after_quote) if part.strip()]
            if after_parts and not re.search(r"\d+\s*万|项目负责人|课题负责人", after_parts[0]):
                source = after_parts[0]
        if "课题" in text and "——" in text:
            topic = text.split("——", 1)[1]
            topic = re.split(r"\(|（", topic, maxsplit=1)[0].strip(" ，,")
            if topic and topic not in title:
                title = f"{title}：{topic}"
    else:
        parts = [part.strip() for part in re.split(r"，|,", text) if part.strip()]
        if hosted:
            title = parts[0] if parts else text
            if len(parts) > 1:
                source = re.sub(r"\s*\d+(?:\.\d+)?\s*万.*$", "", parts[1])
                source = re.sub(r"(项目负责人|课题负责人|参与人|参加人).*$", "", source).strip(" ，,")
        else:
            source = parts[0] if parts else ""
            title = parts[2] if len(parts) > 2 else (parts[1] if len(parts) > 1 else text)
    title = re.sub(r"（[^（）]*(?:万|项目负责人|参与人|参加人|课题负责人|项目骨干)[^（）]*）", "", title)
    title = re.sub(r"\([^()]*(?:万|项目负责人|参与人|参加人|课题负责人|项目骨干)[^()]*\)", "", title)
    return title.strip(" ，,。"), source.strip(" ，,。")


def project_rows(paragraphs: list[str]) -> list[list]:
    rows = []
    order = 1
    for hosted, section_name in [(True, "主持项目"), (False, "参与项目")]:
        for line in get_section(paragraphs, section_name):
            if not re.search(r"\d{4}", line):
                continue
            start_date, end_date = parse_date_range(line)
            title, source = project_title_and_source(line, hosted)
            number = extract_project_number(line)
            role = project_role(line, hosted)
            description = "；".join(part for part in [
                f"角色：{role}",
                f"来源：{source}" if source else "",
                f"项目编号：{number}" if number else "",
                "根据魏雨泉简历（2026.7）整理，公开页仅展示必要信息。",
            ] if part)
            rows.append([
                title[:240],
                number,
                source[:160],
                "魏雨泉" if hosted else "团队参与",
                project_status(end_date),
                start_date,
                end_date,
                extract_amount(line),
                "公开",
                description,
                order if order <= 6 else 0,
            ])
            order += 1
    return rows


def patent_rows(paragraphs: list[str]) -> list[list]:
    rows: OrderedDict[str, list] = OrderedDict()
    for line in get_section(paragraphs, "申请专利"):
        if "专利" not in line or "，" not in line:
            continue
        inventors, rest = line.split(".", 1) if ". " in line else line.split("，", 1)
        inventors = inventors.strip(" .，")
        number_match = re.search(r"(?:专利申请号|申请号|专利号)：?([^），)]+)", rest)
        patent_number = number_match.group(1).replace(" ", "") if number_match else ""
        if not patent_number:
            fallback = re.search(r"(?:CN|ZL)?\s*\d{8,}(?:\.\d+)?[A-Z]?", rest)
            patent_number = fallback.group(0).replace(" ", "") if fallback else ""
        title = re.sub(r"（[^（）]*(?:专利申请号|申请号|专利号|授权公告号|公布号)[^（）]*）", "", rest)
        title = re.sub(r"\([^()]*(?:专利申请号|申请号|专利号|授权公告号|公布号)[^()]*\)", "", title)
        title = re.sub(r"，?\s*(发明专利|实用新型专利).*", "", title).strip(" ，。")
        status = "已授权" if ("已授权" in line or "授权公告号" in line or patent_number.startswith("ZL")) else "申请中"
        if "实用新型" in line:
            status += "；实用新型"
        elif "发明专利" in line:
            status += "；发明专利"
        key = patent_number or title
        if key not in rows:
            rows[key] = [title[:240], patent_number, status, inventors, "", "", "", "公开", len(rows) + 1 if len(rows) < 6 else 0]
    return list(rows.values())


def award_rows(paragraphs: list[str]) -> list[list]:
    rows = []
    for index, line in enumerate(get_section(paragraphs, "科研获奖"), start=1):
        if not re.match(r"\d{4}年", line):
            continue
        date = year_month_date(line)
        parts = [part.strip() for part in re.split(r"，|,", line) if part.strip()]
        title = parts[1] if len(parts) > 1 else line
        rank_match = re.search(r"排名[^）)]+", line)
        rank = rank_match.group(0) if rank_match else ""
        level_parts: list[str] = []
        participant_parts: list[str] = []
        if len(parts) >= 3:
            level_parts.append(parts[2])
        if len(parts) >= 4:
            if "奖" in parts[3] or "成果" in parts[3]:
                level_parts.append(parts[3])
                participant_parts = parts[4:]
            else:
                participant_parts = parts[3:]
        level = "，".join(level_parts)
        level = re.sub(r"（排名[^）]+）", "", level).strip("，, ")
        participants = "，".join(participant_parts)
        participants = re.sub(r"（排名[^）]+）", "", participants).strip("，, ")
        rows.append([title[:240], level[:120], date, participants, "", "", "公开", rank, index if index <= 4 else 0])
    return rows


def clean_title(title: str) -> str:
    title = re.sub(r"^/\.\s*", "", title)
    title = re.sub(r"\s+", " ", title)
    return title.strip(" .")


def parse_publication(line: str, section_name: str) -> dict | None:
    text = re.sub(r"^\d+\.\s*", "", line).strip()
    year_match = re.search(r"\b(20\d{2}|19\d{2})(?:[a-d])?\b", text)
    if not year_match:
        return None
    year = int(year_match.group(1))
    authors = text[: year_match.start()].strip(" .")
    rest = text[year_match.end() :].strip(" .")
    title = ""
    journal = ""
    if ". " in rest:
        title, journal = rest.split(". ", 1)
    elif ". " not in rest and "。" in rest:
        title, journal = rest.split("。", 1)
    else:
        chunks = rest.split(".")
        if len(chunks) >= 2:
            title = chunks[0]
            journal = ".".join(chunks[1:])
        else:
            title = rest
    doi_match = re.search(r"(?i)\bdoi[:：]?\s*(10\.\S+)", text)
    doi = doi_match.group(1).rstrip(" .。") if doi_match else ""
    journal = re.sub(r"(?i)\bdoi[:：]?\s*10\.\S+", "", journal).strip(" .。")
    return {
        "authors": authors,
        "year": year,
        "title": clean_title(title),
        "journal_raw": journal,
        "doi": doi,
        "section": section_name,
        "source_line": text,
    }


def split_journal_meta(raw: str) -> tuple[str, str, str, str]:
    text = raw.strip(" .。")
    pages = ""
    volume = ""
    issue = ""
    page_match = re.search(r"(\d+\s*[-–]\s*\d+|\d{5,7}|e\d+)\.?$", text)
    if page_match:
        pages = page_match.group(1).replace(" ", "")
        text = text[: page_match.start()].strip(" ,，.")
    issue_match = re.search(r"(\d+)\s*\(([^)]+)\)\s*$", text)
    if issue_match:
        volume, issue = issue_match.group(1), issue_match.group(2)
        text = text[: issue_match.start()].strip(" ,，.")
    else:
        tail = re.search(r",\s*([\d]+)\s*$", text)
        if tail:
            volume = tail.group(1)
            text = text[: tail.start()].strip(" ,，.")
    journal = text.strip(" ,，.")
    return journal, volume, issue, pages


def load_enrichment_cache() -> dict[str, dict]:
    cache: dict[str, dict] = {}
    for filename in [
        "publication_enrichment_cache.json",
        "publication_europepmc_cache.json",
        "publication_semantic_scholar_cache.json",
    ]:
        path = OUTPUT_DIR / filename
        if not path.exists():
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            continue
        for key, value in data.items():
            if isinstance(value, dict) and value.get("matched"):
                cache[key] = value
    return cache


def find_enrichment(pub: dict, cache: dict[str, dict]) -> dict:
    keys = [
        f"{pub.get('title')}::{pub.get('year')}",
        f"{pub.get('doi')}::{pub.get('year')}" if pub.get("doi") else "",
    ]
    for key in keys:
        if key and key in cache:
            return cache[key]
    normalized_title = re.sub(r"\W+", "", pub.get("title", "")).lower()
    for key, value in cache.items():
        if re.sub(r"\W+", "", str(value.get("title", ""))).lower() == normalized_title:
            return value
    return {}


def publication_rows(paragraphs: list[str]) -> list[list]:
    parsed: list[dict] = []
    for section_name in ["第一作者或通讯作者发表文章 （非农大第一单位标红）", "参与发表文章"]:
        for line in get_section(paragraphs, section_name):
            pub = parse_publication(line, section_name)
            if pub and pub["title"]:
                parsed.append(pub)
    seen = set()
    unique = []
    for pub in parsed:
        key = (re.sub(r"\W+", "", pub["title"]).lower(), pub["year"])
        if key in seen:
            continue
        seen.add(key)
        unique.append(pub)

    cache = load_enrichment_cache()
    rows = []
    for index, pub in enumerate(sorted(unique, key=lambda item: (-item["year"], item["title"])), start=1):
        enrichment = find_enrichment(pub, cache)
        journal, volume, issue, pages = split_journal_meta(pub["journal_raw"])
        doi = pub["doi"] or str(enrichment.get("doi") or "").replace("https://doi.org/", "")
        abstract = str(enrichment.get("abstract") or "")
        journal = str(enrichment.get("journal") or journal)
        volume = str(enrichment.get("volume") or volume)
        issue = str(enrichment.get("issue") or issue)
        pages = str(enrichment.get("pages") or pages)
        # 首页只放少量最新一作/通讯论文；其余保留在成果页。
        home_order = index if pub["section"].startswith("第一作者") and index <= 8 else 0
        rows.append([
            pub["source_line"],
            abstract,
            "公开",
            "",
            home_order,
        ])
    return rows


def student_rows(tables: list[list[list[str]]]) -> list[list]:
    rows = []
    for table in tables:
        if len(table) < 3:
            continue
        title = table[0][0]
        headers = table[1]
        for raw in table[2:]:
            if len(raw) < 5 or not raw[1] or raw[1] == "姓名":
                continue
            name = raw[1]
            training = raw[2] if len(raw) > 2 else ""
            grade = raw[3] if len(raw) > 3 else ""
            degree = "博士" if "博士" in title or "博士" in training else ("硕士" if "硕士" in title or "硕" in training else "")
            graduation = raw[4] if "毕业" in title and len(raw) > 4 else ""
            topic = raw[5] if "毕业" in title and len(raw) > 5 else (raw[4] if len(raw) > 4 else "")
            advisor = "李季" if "李季" in training else ("许艇" if "许艇" in training else ("乔玉辉" if "乔玉辉" in training else "魏雨泉"))
            rows.append([name, degree, grade, training, graduation, topic, advisor, title.replace("：", "")])
    return rows


def write_workbook(filename: str, title: str, headers: list[str], rows: list[list], notes: list[str], widths: list[int]) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    info = workbook.active
    info.title = "导入说明"
    info["A1"] = title
    info["A1"].font = Font(bold=True, size=16, color="1F3D2B")
    info["A2"] = "填写说明"
    info["A2"].fill = TITLE_FILL
    info["A2"].font = TITLE_FONT
    for index, note in enumerate(notes, start=3):
        info[f"A{index}"] = f"{index - 2}. {note}"
        info[f"A{index}"].alignment = Alignment(wrap_text=True, vertical="top")
    info.column_dimensions["A"].width = 110

    data = workbook.create_sheet("导入数据")
    data.append(headers)
    for row in rows:
        data.append(row)
    for column, width in enumerate(widths, start=1):
        data.column_dimensions[data.cell(row=1, column=column).column_letter].width = width
    for cell in data[1]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = CELL_BORDER
    for row in data.iter_rows(min_row=2, max_row=max(2, data.max_row), max_col=len(headers)):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = CELL_BORDER
    data.freeze_panes = "A2"
    data.auto_filter.ref = data.dimensions

    audit = workbook.create_sheet("整理记录")
    audit.append(["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    audit.append(["数据来源", str(find_source())])
    audit.append(["记录数", len(rows)])
    audit.append(["说明", "本表按当前 LabHub 门户内容导入字段整理，附件和图片需在单条内容编辑页补充。"])
    for row in audit.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    audit.column_dimensions["A"].width = 16
    audit.column_dimensions["B"].width = 100

    path = OUTPUT_DIR / filename
    workbook.save(path)
    return path


def write_readme(paths: list[Path], counts: dict[str, int]) -> Path:
    readme = OUTPUT_DIR / "README-导入说明.md"
    lines = [
        "# 魏雨泉简历导入文件说明",
        "",
        f"数据来源：`{find_source()}`",
        "",
        "这版重新按 LabHub 后台模板整理：第二张表统一为 `导入数据`，使用中文表头，并增加 `整理记录` 便于核对。",
        "",
        "## 文件",
    ]
    for path in paths:
        lines.append(f"- `{path.name}`")
    lines += [
        "",
        "## 记录数",
        f"- 团队成员：{counts['members']} 条",
        f"- 科研项目：{counts['projects']} 条",
        f"- 科研获奖：{counts['awards']} 条",
        f"- 专利成果：{counts['patents']} 条",
        f"- 论文成果：{counts['publications']} 条",
        f"- 学生清单：{counts['students']} 条（当前作为核对清单，学生档案暂未提供批量导入接口）",
        "",
        "## 建议导入顺序",
        "团队成员 -> 科研项目 -> 专利成果 -> 论文成果 -> 获奖成果。",
        "",
        "论文、专利、获奖的 PDF/图片附件不在本次 Excel 中直接嵌入，建议导入文字信息后在单条编辑页补充。",
    ]
    readme.write_text("\n".join(lines), encoding="utf-8")
    return readme


def main() -> None:
    source = find_source()
    paragraphs, tables = load_docx(source)

    members = member_rows(paragraphs)
    projects = project_rows(paragraphs)
    patents = patent_rows(paragraphs)
    awards = award_rows(paragraphs)
    publications = publication_rows(paragraphs)
    students = student_rows(tables)

    paths = [
        write_workbook(
            "01-团队成员-魏雨泉.xlsx",
            "团队成员批量导入文件",
            MEMBER_HEADERS,
            members,
            ["导入位置：门户内容 -> 团队成员。头像需在成员编辑页单独上传。", "展示排序大于 0 时会在公开网站按顺序展示。"],
            [14, 22, 30, 34, 24, 90, 12],
        ),
        write_workbook(
            "02-科研项目-魏雨泉.xlsx",
            "科研项目批量导入文件",
            PROJECT_HEADERS,
            projects,
            ["导入位置：门户内容 -> 科研项目。", "项目说明保留角色、来源和编号；经费字段保留在后台，前台可按页面逻辑控制是否展示。"],
            [42, 24, 34, 16, 12, 14, 14, 10, 14, 64, 12],
        ),
        write_workbook(
            "03-获奖成果-魏雨泉.xlsx",
            "获奖成果批量导入文件",
            AWARD_HEADERS,
            awards,
            ["导入位置：门户内容 -> 获奖成果。", "仅整理科研和技术成果类奖项，教学获奖、学生竞赛指导获奖未放入成果库。"],
            [44, 30, 14, 58, 20, 20, 14, 24, 12],
        ),
        write_workbook(
            "04-专利成果-魏雨泉.xlsx",
            "专利成果批量导入文件",
            PATENT_HEADERS,
            patents,
            ["导入位置：门户内容 -> 专利成果。", "重复专利号已自动合并；专利 PDF 可后续在单条编辑页补充。"],
            [46, 24, 16, 52, 20, 14, 14, 14, 12],
        ),
        write_workbook(
            "05-论文成果-魏雨泉.xlsx",
            "论文成果批量导入文件",
            PUBLICATION_HEADERS,
            publications,
            ["导入位置：门户内容 -> 论文成果。", "只保留 GB/T 7714-2025 格式引文，系统自动拆分作者、题目、期刊、年份、卷期页和 DOI。", "优先按 DOI 更新；无 DOI 时按论文题目和年份匹配。", "摘要来自已有外部检索缓存；未检索到真实摘要的条目留空。"],
            [82, 80, 14, 20, 12],
        ),
        write_workbook(
            "06-指导学生清单-魏雨泉.xlsx",
            "指导学生清单",
            STUDENT_HEADERS,
            students,
            ["当前系统暂无学生档案批量导入入口，本表先作为核对和后续导入开发依据。", "导师列为根据培养情况自动推断，正式导入前建议人工确认。"],
            [14, 12, 12, 36, 14, 64, 16, 18],
        ),
    ]
    counts = {
        "members": len(members),
        "projects": len(projects),
        "awards": len(awards),
        "patents": len(patents),
        "publications": len(publications),
        "students": len(students),
    }
    readme = write_readme(paths, counts)
    for path in paths + [readme]:
        print(path)


if __name__ == "__main__":
    main()
