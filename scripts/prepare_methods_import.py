import re
import shutil
import zipfile
import posixpath
from copy import deepcopy
from pathlib import Path
from tempfile import NamedTemporaryFile
from xml.etree import ElementTree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


SOURCE = Path(r"C:/Users/baoju/Desktop/团队实验方法汇总0709.docx")
OUTPUT_DIR = Path(__file__).resolve().parents[1] / "outputs" / "internal-documents" / "methods-import-0709"
DOCS_DIR = OUTPUT_DIR / "files"
XLSX_PATH = OUTPUT_DIR / "团队实验方法批量导入清单.xlsx"
ZIP_PATH = OUTPUT_DIR / "团队实验方法批量导入文件包.zip"


def iter_blocks(document):
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def text_of(block):
    if isinstance(block, Paragraph):
        return block.text.strip()
    if isinstance(block, Table):
        rows = []
        for row in block.rows:
            rows.append("；".join(cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()))
        return " ".join(rows).strip()
    return ""


def paragraph_has_image(block):
    if not isinstance(block, Paragraph):
        return False
    return bool(block._p.xpath(".//*[local-name()='blip']"))


def is_heading2(block):
    return isinstance(block, Paragraph) and block.style and block.style.name == "Heading 2" and block.text.strip()


def is_heading1(block):
    return isinstance(block, Paragraph) and block.style and block.style.name == "Heading 1" and block.text.strip()


def is_section_heading_text(text):
    return bool(re.match(r"^第[一二三四五六七八九十]+部分\\s*", text.strip()))


def clean_title(text):
    text = re.sub(r"^\s*\d+(?:\.\d+)*\s*", "", text.strip())
    text = text.replace("检测方法", "").strip(" ：:，,")
    return text or "实验方法"


def safe_filename(index, title):
    name = re.sub(r'[\\/:*?"<>|]+', "", title)
    name = re.sub(r"\s+", "", name)
    if len(name) > 36:
        name = name[:36]
    return f"{index:02d}-{name}.docx"


def natural_section_label(section):
    text = re.sub(r"^第[一二三四五六七八九十]+部分\s*", "", section or "").strip()
    if not text:
        return ""
    if text.endswith(("实验", "培养", "测定", "分析", "使用")):
        return f"{text}方法"
    return text


def first_description(blocks, section=""):
    location = ""
    equipment = ""
    for block in blocks:
        text = text_of(block)
        if not text:
            continue
        if text.startswith("实验地点") and not location:
            location = text.replace("实验地点：", "").replace("实验地点:", "").strip()
        elif text.startswith("实验器材") and not equipment:
            equipment = text.replace("实验器材：", "").replace("实验器材:", "").strip()
        if location and equipment:
            break
    parts = []
    section_label = natural_section_label(section)
    if section_label:
        parts.append(section_label)
    if location:
        parts.append(f"地点：{location}")
    if equipment:
        parts.append(f"器材：{equipment}")
    desc = "；".join(parts) or "团队实验方法汇总。"
    return re.sub(r"\s+", " ", desc).strip()[:80]


def copy_paragraph(target, paragraph):
    new_p = target.add_paragraph()
    new_p._p.getparent().replace(new_p._p, deepcopy(paragraph._p))


def copy_table(target, table):
    new_p = target.add_paragraph()
    new_p._p.addnext(deepcopy(table._tbl))
    p = new_p._p
    p.getparent().remove(p)


def build_method_doc(title, blocks, output_path):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()

    for block in blocks:
        if isinstance(block, Paragraph):
            if (block.text.strip() or paragraph_has_image(block)) and not is_section_heading_text(block.text):
                copy_paragraph(doc, block)
        elif isinstance(block, Table):
            copy_table(doc, block)

    doc.save(output_path)
    patch_media_relationships(output_path, SOURCE)


def patch_media_relationships(docx_path, source_docx):
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    doc_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rel_attr_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    content_type_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    image_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"

    ElementTree.register_namespace("", rel_ns)
    ElementTree.register_namespace("w", doc_ns)
    ElementTree.register_namespace("r", rel_attr_ns)

    with zipfile.ZipFile(docx_path, "r") as target_zip, zipfile.ZipFile(source_docx, "r") as source_zip:
        document_xml = target_zip.read("word/document.xml")
        document_root = ElementTree.fromstring(document_xml)
        used_rel_ids = {
            value
            for elem in document_root.iter()
            for key, value in elem.attrib.items()
            if key in {f"{{{rel_attr_ns}}}embed", f"{{{rel_attr_ns}}}link"}
        }
        if not used_rel_ids:
            return

        source_rels_root = ElementTree.fromstring(source_zip.read("word/_rels/document.xml.rels"))
        source_rels = {
            rel.attrib.get("Id"): rel.attrib
            for rel in source_rels_root.findall(f"{{{rel_ns}}}Relationship")
            if rel.attrib.get("Type") == image_rel_type
        }

        target_rels_root = ElementTree.fromstring(target_zip.read("word/_rels/document.xml.rels"))
        existing_ids = {rel.attrib.get("Id") for rel in target_rels_root.findall(f"{{{rel_ns}}}Relationship")}

        content_types_root = ElementTree.fromstring(target_zip.read("[Content_Types].xml"))
        existing_defaults = {
            default.attrib.get("Extension")
            for default in content_types_root.findall(f"{{{content_type_ns}}}Default")
        }

        remap = {}
        media_to_copy = []
        next_index = 1
        for old_id in sorted(used_rel_ids):
            source_rel = source_rels.get(old_id)
            if not source_rel:
                continue
            target = source_rel.get("Target", "")
            source_media_path = posixpath.normpath(posixpath.join("word", target))
            if source_media_path not in source_zip.namelist():
                continue
            while f"rIdImportedImage{next_index}" in existing_ids:
                next_index += 1
            new_id = f"rIdImportedImage{next_index}"
            existing_ids.add(new_id)
            next_index += 1

            media_name = Path(source_media_path).name
            new_target = f"media/{old_id}-{media_name}"
            new_media_path = f"word/{new_target}"
            remap[old_id] = new_id
            media_to_copy.append((source_media_path, new_media_path))
            ElementTree.SubElement(
                target_rels_root,
                f"{{{rel_ns}}}Relationship",
                {"Id": new_id, "Type": image_rel_type, "Target": new_target},
            )

            extension = Path(media_name).suffix.lower().lstrip(".")
            content_type = {
                "emf": "image/x-emf",
                "gif": "image/gif",
                "jpeg": "image/jpeg",
                "jpg": "image/jpeg",
                "png": "image/png",
                "svg": "image/svg+xml",
                "webp": "image/webp",
                "wmf": "image/x-wmf",
            }.get(extension)
            if extension and content_type and extension not in existing_defaults:
                ElementTree.SubElement(
                    content_types_root,
                    f"{{{content_type_ns}}}Default",
                    {"Extension": extension, "ContentType": content_type},
                )
                existing_defaults.add(extension)

        if not remap:
            return

        for elem in document_root.iter():
            for attr in (f"{{{rel_attr_ns}}}embed", f"{{{rel_attr_ns}}}link"):
                if elem.attrib.get(attr) in remap:
                    elem.set(attr, remap[elem.attrib[attr]])

        with NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_path = Path(temp_file.name)

        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as output_zip:
            for info in target_zip.infolist():
                if info.filename in {"word/document.xml", "word/_rels/document.xml.rels", "[Content_Types].xml"}:
                    continue
                output_zip.writestr(info, target_zip.read(info.filename))
            output_zip.writestr("word/document.xml", ElementTree.tostring(document_root, encoding="utf-8", xml_declaration=True))
            output_zip.writestr("word/_rels/document.xml.rels", ElementTree.tostring(target_rels_root, encoding="utf-8", xml_declaration=True))
            output_zip.writestr("[Content_Types].xml", ElementTree.tostring(content_types_root, encoding="utf-8", xml_declaration=True))
            for source_media_path, new_media_path in media_to_copy:
                output_zip.writestr(new_media_path, source_zip.read(source_media_path))

    shutil.move(temp_path, docx_path)


def extract_methods(source):
    source_doc = Document(source)
    methods = []
    current = None
    current_section = ""
    for block in iter_blocks(source_doc):
        if is_heading1(block):
            if current:
                methods.append(current)
                current = None
            current_section = re.sub(r"\s+", " ", block.text.strip())
            continue
        if is_heading2(block):
            if current:
                methods.append(current)
            current = {"title": clean_title(block.text), "section": current_section, "blocks": []}
            continue
        if current is not None:
            text = text_of(block)
            if text or isinstance(block, Table) or paragraph_has_image(block):
                current["blocks"].append(block)
    if current:
        methods.append(current)
    return methods


def style_sheet(sheet):
    sheet.freeze_panes = "A2"
    header_fill = PatternFill("solid", fgColor="00873C")
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def build_import_xlsx(rows):
    wb = Workbook()
    guide = wb.active
    guide.title = "导入说明"
    guide["A1"] = "团队实验方法批量导入说明"
    guide["A1"].font = Font(size=16, bold=True, color="1F3D2B")
    guide["A3"] = "使用方式"
    guide["B3"] = "在内部资料库点击“批量导入”，上传本清单和同目录下的 zip 文件包。"
    guide["A4"] = "分类"
    guide["B4"] = "本批资料统一归入“实验方法”。"
    guide["A5"] = "权限"
    guide["B5"] = "默认“组内成员可见”，允许下载。"
    guide.column_dimensions["A"].width = 16
    guide.column_dimensions["B"].width = 86
    for row in range(3, 6):
        guide[f"A{row}"].font = Font(bold=True, color="1F3D2B")
        guide[f"A{row}"].fill = PatternFill("solid", fgColor="EAF5EE")
        guide[f"B{row}"].alignment = Alignment(wrap_text=True)

    data = wb.create_sheet("导入数据")
    headers = ["资料标题", "资料分类", "可见范围", "资料说明", "文件名", "允许下载"]
    data.append(headers)
    for row in rows:
        data.append(row)
    style_sheet(data)
    widths = [34, 18, 18, 58, 42, 12]
    for index, width in enumerate(widths, start=1):
        data.column_dimensions[get_column_letter(index)].width = width
    for index in range(2, len(rows) + 2):
        data.row_dimensions[index].height = 48

    wb.save(XLSX_PATH)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    methods = extract_methods(SOURCE)
    rows = []
    filenames = []
    for index, method in enumerate(methods, start=1):
        title = method["title"]
        filename = safe_filename(index, title)
        output_path = DOCS_DIR / filename
        build_method_doc(title, method["blocks"], output_path)
        rows.append([title, "实验方法", "组内成员可见", first_description(method["blocks"], method.get("section", "")), filename, "是"])
        filenames.append(filename)

    build_import_xlsx(rows)
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as archive:
        for filename in filenames:
            archive.write(DOCS_DIR / filename, arcname=filename)

    print(f"methods={len(methods)}")
    print(XLSX_PATH)
    print(ZIP_PATH)


if __name__ == "__main__":
    main()
