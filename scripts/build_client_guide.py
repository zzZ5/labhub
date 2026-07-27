from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "中农雨磷课题组网站使用与管理说明.docx"
LOGO = ROOT / "frontend" / "public" / "site-icon.png"

GREEN = "00873C"
DEEP_GREEN = "1F3D2B"
PALE_GREEN = "EAF5EE"
WARM_WHITE = "F8F7F2"
LIGHT_GRAY = "F5F7F6"
BORDER = "D9E1DC"
TEXT = "2F3437"
MUTED = "6B7280"
SOIL_GOLD = "A2622A"
BLUE = "00495E"
WHITE = "FFFFFF"
MANUAL_NUMBER_COUNTERS = {}


def set_run_font(run, size=11, bold=False, color=TEXT, italic=False, east_asia="SimSun"):
    run.font.name = "Times New Roman"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size=6):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent))
    tblInd.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    pPr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    pPr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    lvl.append(pPr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def add_body(doc, text, bold_lead=None, color=TEXT, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=DEEP_GREEN)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color)
    return p


def add_bullet(doc, text, num_id, bold_lead=None):
    p = add_body(doc, text, bold_lead=bold_lead, after=4)
    apply_numbering(p, num_id)
    return p


def add_numbered(doc, text, num_id, bold_lead=None):
    number = MANUAL_NUMBER_COUNTERS.get(num_id, 0) + 1
    MANUAL_NUMBER_COUNTERS[num_id] = number
    p = add_body(doc, f"{number}.  {text}", bold_lead=bold_lead, after=4)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_callout(doc, title, text, tone="green"):
    fills = {"green": PALE_GREEN, "gray": LIGHT_GRAY, "gold": "F6F0E7", "blue": "EAF2F4"}
    accents = {"green": GREEN, "gray": MUTED, "gold": SOIL_GOLD, "blue": BLUE}
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fills[tone])
    set_cell_border(cell, color=fills[tone], size=2)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=accents[tone])
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers, rows, widths, header_fill=PALE_GREEN, compact=False):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=9.5 if compact else 10, bold=True, color=DEEP_GREEN)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for col_index, (cell, value) in enumerate(zip(cells, values)):
            if row_index % 2 == 1:
                shade_cell(cell, "FBFCFB")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=9.5 if compact else 10, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_checklist(doc, rows):
    return add_table(doc, ["核对", "系统检查项", "建议结果"], [["□", item, result] for item, result in rows], [700, 6100, 2560], header_fill=LIGHT_GRAY, compact=True)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal_fonts = normal._element.rPr.rFonts
    normal_fonts.set(qn("w:ascii"), "Times New Roman")
    normal_fonts.set(qn("w:hAnsi"), "Times New Roman")
    normal_fonts.set(qn("w:cs"), "Times New Roman")
    normal_fonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, GREEN, 18, 10),
        "Heading 2": (13, GREEN, 14, 7),
        "Heading 3": (12, DEEP_GREEN, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style_fonts = style._element.rPr.rFonts
        style_fonts.set(qn("w:ascii"), "Times New Roman")
        style_fonts.set(qn("w:hAnsi"), "Times New Roman")
        style_fonts.set(qn("w:cs"), "Times New Roman")
        style_fonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("中农雨磷")
    set_run_font(left, size=9, bold=True, color=DEEP_GREEN)
    right = p.add_run("  |  网站使用与管理说明")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        run = p.add_run()
        shape = run.add_picture(str(LOGO), width=Inches(0.78))
        doc_pr = shape._inline.docPr
        doc_pr.set("descr", "中农雨磷课题组标识")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("中国农业大学资源与环境学院")
    set_run_font(r, size=11, bold=True, color=GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("中农雨磷课题组")
    set_run_font(r, size=24, bold=True, color=DEEP_GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("门户网站与内部科研管理平台")
    set_run_font(r, size=18, bold=True, color=GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("使用与管理说明")
    set_run_font(r, size=14, color=MUTED)

    add_callout(
        doc,
        "阅读说明",
        "本说明书采用日常操作语言编写，不要求具备编程知识。按照本文操作，即可完成官网内容维护、账号管理、资料归档、仪器信息维护和学生材料管理。",
        tone="green",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("版本：使用说明 V1.0")
    set_run_font(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("编制日期：2026 年 7 月 22 日")
    set_run_font(r, size=10, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    r = p.add_run("适用地点：中国农业大学西校区")
    set_run_font(r, size=10, color=MUTED)
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    bullet_num = add_numbering(doc, "bullet")
    MANUAL_NUMBER_COUNTERS.clear()
    decimal_num = 0
    add_cover(doc)

    add_heading(doc, "一、这套网站是什么", 1)
    add_body(doc, "LabHub 是为中农雨磷课题组建设的一套统一网站，包含面向校内外访客的公开门户，以及仅供组内成员登录使用的内部科研管理平台。两部分采用同一品牌形象，但承担不同任务。")
    add_table(
        doc,
        ["组成部分", "主要对象", "核心作用"],
        [
            ["公开门户", "学生、同行、合作单位及社会访客", "展示研究方向、团队成员、科研成果和新闻活动，建立规范的课题组对外形象。"],
            ["内部平台", "已审核的课题组成员", "集中管理内部资料、仪器信息、学生档案、账号权限和门户内容。"],
        ],
        [1700, 2600, 5060],
    )
    add_callout(doc, "核心原则", "公开信息与内部资料分开管理；学校身份与系统权限分开设置；内容由业务人员在网页后台维护，无需修改程序。", tone="gold")

    add_heading(doc, "二、使用者和权限", 1)
    add_body(doc, "系统不会仅凭“导师”“博士生”等学校身份自动赋予管理权限。学校身份用于说明成员在课题组中的身份，系统权限用于决定其能否管理网站、资料、仪器或账号，两者可以自由组合。")
    add_table(
        doc,
        ["使用者", "默认可做的事情", "需要额外权限的事情"],
        [
            ["校外访客", "浏览公开官网内容", "不能进入内部平台"],
            ["已审核组内成员", "进入工作台、查阅内部资料和仪器信息、查看学生档案、维护个人资料", "不能默认修改门户、设备或他人资料"],
            ["网站编辑", "维护官网首页、横幅、研究方向、成员、新闻和成果", "不能默认管理账号"],
            ["资料管理员", "整理、编辑和删除内部资料", "不能默认维护门户内容"],
            ["仪器管理员", "新增、编辑和删除仪器信息", "不能默认管理账号"],
            ["系统管理员", "账号、权限和全部业务模块的管理", "应仅授予少数负责人"],
        ],
        [1500, 4300, 3560],
        compact=True,
    )
    add_bullet(doc, "学校身份包括：硕博导师、博士后、博士生、硕士生、本科生和其他成员。", bullet_num)
    add_bullet(doc, "成员状态分为“在组”和“已毕业/离组”；状态变化不会自动删除账号和历史资料。", bullet_num)
    add_bullet(doc, "账号可以使用账号名或邮箱登录；无邮箱时也可以创建账号。", bullet_num)
    add_bullet(doc, "新账号需要审核通过后才能进入内部平台。", bullet_num)

    add_heading(doc, "三、公开门户包含哪些内容", 1)
    add_body(doc, "公开门户是课题组对外展示窗口，整体采用中国农业大学生命绿和克制的学术风格，支持电脑、平板和手机访问。")
    add_table(
        doc,
        ["页面", "展示内容", "维护入口"],
        [
            ["首页", "轮播横幅、课题组简介、研究方向、科研成果、团队成员、新闻活动、联系方式", "内部平台 → 门户内容"],
            ["研究方向", "方向列表、封面图、摘要和图文详情", "门户内容 → 研究方向"],
            ["团队成员", "头像、姓名、身份头衔、研究方向和个人简介", "门户内容 → 团队成员"],
            ["科研成果", "论文、科研项目、专利和获奖成果；支持检索、筛选、分页和详情", "门户内容 → 各成果栏目"],
            ["新闻活动", "分类列表、封面、摘要、图文正文、发布时间和浏览次数", "门户内容 → 新闻活动"],
            ["页脚", "课题组简介、地址和学校/学院等外部链接", "门户内容 → 页脚设置"],
        ],
        [1500, 5060, 2800],
        compact=True,
    )
    add_callout(doc, "展示顺序", "研究方向、团队成员和成果使用“排序”控制展示。排序为 0 时通常不在首页重点展示；大于 0 时数字越小越靠前。首页会限制展示数量，完整内容仍可在列表页查看。", tone="blue")

    add_heading(doc, "四、内部平台包含哪些内容", 1)
    add_heading(doc, "4.1 工作台", 2)
    add_body(doc, "登录后首先进入工作台。工作台用紧凑卡片显示内部资料数量、仪器数量和学生档案数量，并列出近期资料、需要关注的设备和最近更新的学生归档，帮助成员快速进入常用功能。")
    add_heading(doc, "4.2 内部资料库", 2)
    add_bullet(doc, "支持分类、搜索、排序和分页，资料多时仍能快速查找。", bullet_num)
    add_bullet(doc, "支持上传文件、批量导入和添加视频链接；文件与视频链接可以同时存在。", bullet_num)
    add_bullet(doc, "PDF、图片和视频可直接查看；Word、PPT 等文件会尽量转换为网页可预览格式。", bullet_num)
    add_bullet(doc, "普通成员主要维护自己上传的资料；管理员和资料管理员承担整体整理工作。", bullet_num)
    add_heading(doc, "4.3 仪器平台", 2)
    add_bullet(doc, "展示仪器名称、型号、状态、详细位置、图片、负责人和使用说明。", bullet_num)
    add_bullet(doc, "支持搜索、排序、分页、新建、编辑、删除和 Excel 批量导入。", bullet_num)
    add_bullet(doc, "仪器使用在线下登记，系统不提供预约功能，也不设置复杂设备编号。", bullet_num)
    add_heading(doc, "4.4 学生档案", 2)
    add_bullet(doc, "一个学生账号最多对应一个学生档案；导师不需要建立学生档案。", bullet_num)
    add_bullet(doc, "学生可以关联一位主导师和多位导师。", bullet_num)
    add_bullet(doc, "档案保存姓名、学位类型、年级、研究题目、研究方向、入学/毕业日期和毕业去向。", bullet_num)
    add_bullet(doc, "归档资料支持单个上传、拖拽批量上传、在线预览、下载和后续编辑。", bullet_num)
    add_bullet(doc, "组内已审核成员可以查阅学生档案；资料删除等敏感操作仍受权限限制。", bullet_num)
    add_heading(doc, "4.5 成员与账号管理", 2)
    add_bullet(doc, "系统管理员可新建账号、批量导入账号、审核成员、调整学校身份和系统权限。", bullet_num)
    add_bullet(doc, "学生身份账号可一键生成学生档案，系统会阻止同一账号重复建档。", bullet_num)
    add_bullet(doc, "成员毕业或离组后建议先修改成员状态，再根据实际需要决定是否停用登录。", bullet_num)

    add_heading(doc, "五、门户内容如何维护", 1)
    add_body(doc, "拥有“网站编辑”或“系统管理员”权限的人员，可在内部平台进入“门户内容”。这里是公开官网的主要内容来源，保存后公开网站会读取最新内容。")
    add_table(
        doc,
        ["后台栏目", "可维护内容", "前台影响"],
        [
            ["站点首页", "实验室名称、归属单位、横幅副标题、课题组简介、横幅切换间隔", "导航、首页横幅和课题组简介"],
            ["页脚设置", "Logo、网站图标、默认横幅、页脚简介、地址、相关链接", "导航品牌、浏览器图标、首页兜底横幅和页脚"],
            ["首页横幅", "标题、副标题、图片、跳转链接、排序、启用状态", "首页照片轮播"],
            ["研究方向", "标题、摘要、图文正文、封面图和排序", "首页、研究列表和详情页"],
            ["团队成员", "姓名、身份头衔、邮箱、研究方向、头像、图文简介和排序", "首页成员、团队列表和个人详情"],
            ["新闻活动", "标题、分类、日期、地点、摘要、图文正文、Word 稿件、封面和发布状态", "新闻列表和详情页"],
            ["各类成果", "论文、项目、专利、获奖的信息、图文说明、附件和首页排序", "首页成果和科研成果页"],
        ],
        [1650, 4950, 2760],
        compact=True,
    )
    add_callout(doc, "重要提醒", "创建内部账号不会自动把该人员公开到团队页面；公开团队成员需在“门户内容 → 团队成员”中单独维护。这样可以避免未准备公开的信息意外出现在官网。", tone="gold")

    add_heading(doc, "六、图文详情如何编辑", 1)
    add_body(doc, "新闻正文、个人简介、研究方向正文、论文摘要、项目说明、专利说明和获奖说明均使用统一的图文编辑器。使用方式与常见文档软件相近。")
    decimal_num += 1
    add_numbered(doc, "先填写标题、摘要、日期等基础信息。", decimal_num)
    add_numbered(doc, "在正文编辑区输入文字，可设置二级/三级标题、粗体、斜体、列表、引用、对齐和链接。", decimal_num)
    add_numbered(doc, "把光标放到需要的位置，点击“插图”选择图片；上传完成后图片会插入当前正文位置。", decimal_num)
    add_numbered(doc, "检查移动端阅读效果和图片位置，再点击保存。", decimal_num)
    add_numbered(doc, "打开公开详情页复核内容，确认无误后再对外发布或分享。", decimal_num)
    add_bullet(doc, "新闻可直接上传 DOCX 稿件，系统会尝试提取文字和内嵌图片，转换完成后再展示。", bullet_num)
    add_bullet(doc, "插图会自动进入门户图片资源，并使用系统的图片压缩机制，避免原图过大拖慢网页。", bullet_num)
    add_bullet(doc, "列表页只显示纯文本摘要，不会显示编辑器中的格式代码；完整图文仅在详情页展示。", bullet_num)
    add_bullet(doc, "详情页会记录浏览次数，并显示发布时间或最近更新时间。", bullet_num)

    add_heading(doc, "七、新闻发布建议流程", 1)
    decimal_num += 1
    add_numbered(doc, "准备标题、简短摘要、活动日期、分类和封面图。", decimal_num)
    add_numbered(doc, "选择“学术交流、组内动态、科研进展、成果荣誉、招生招聘、项目动态”之一。", decimal_num)
    add_numbered(doc, "在正文编辑器中排版，或上传 DOCX 稿件后继续修改。", decimal_num)
    add_numbered(doc, "保存为草稿并预览，检查标题、图片、段落和手机端显示。", decimal_num)
    add_numbered(doc, "将状态改为发布；需要长期保留但不再突出时可归档。", decimal_num)
    add_callout(doc, "图片建议", "封面优先使用横向照片；正文图片应清晰但不必上传相机原始大图。人物合影、田间试验、实验过程和成果照片比通用素材更适合课题组网站。", tone="green")

    add_heading(doc, "八、科研成果管理", 1)
    add_body(doc, "科研成果分为论文、科研项目、专利和获奖成果，四类内容独立管理、独立检索和分页。")
    add_table(
        doc,
        ["成果类型", "建议填写内容", "附件与详情"],
        [
            ["论文", "GB/T 7714—2025 格式引文、摘要、DOI、年份和期刊信息", "可上传公开 PDF；详情页可展示图文摘要和 DOI 跳转"],
            ["科研项目", "项目名称、编号、来源、负责人、周期、状态和适合公开的说明", "公开信息宜简洁，避免上传敏感材料"],
            ["专利", "专利名称、专利号、发明人、申请/授权日期、状态和成果说明", "可上传 PDF，说明支持插图"],
            ["获奖成果", "奖项名称、等级、日期、参与人员和成果说明", "可上传图片及 PDF/图片附件"],
        ],
        [1450, 4650, 3260],
        compact=True,
    )
    add_bullet(doc, "成果列表默认优先按时间从新到旧展示；设置首页排序后可人工控制首页重点成果。", bullet_num)
    add_bullet(doc, "论文可按年份和关键词检索，项目、专利、获奖也支持关键词检索和分页。", bullet_num)
    add_bullet(doc, "批量导入时系统会按 DOI、项目编号、专利号或名称等规则尽量识别重复内容，更新已有记录。", bullet_num)

    add_heading(doc, "九、内部资料库使用说明", 1)
    add_heading(doc, "9.1 上传一份资料", 2)
    decimal_num += 1
    add_numbered(doc, "进入“内部资料”，先选择左侧分类。", decimal_num)
    add_numbered(doc, "点击“上传资料”，当前分类会自动带入，可按需要修改。", decimal_num)
    add_numbered(doc, "填写资料标题和简短说明，可上传文件、填写视频链接，或两者同时提供。", decimal_num)
    add_numbered(doc, "选择是否允许下载，然后保存并等待预览生成完成。", decimal_num)
    add_heading(doc, "9.2 在线预览", 2)
    add_table(
        doc,
        ["资料形式", "查看方式", "说明"],
        [
            ["PDF", "网页内直接阅读", "推荐用于长期归档和稳定展示"],
            ["图片", "网页内直接查看", "适合图表、流程图和照片"],
            ["Word / PPT", "后台转换后在线预览", "转换需要时间；复杂动画或特殊字体可能与原文件略有差异"],
            ["视频链接", "网页内解析播放或跳转", "支持常见 B 站链接，短链接会尝试转换"],
            ["其他文件", "下载查看", "无法转换时会给出明确提示"],
        ],
        [1500, 2500, 5360],
        compact=True,
    )
    add_callout(doc, "预览失败不等于文件丢失", "若 Word 或 PPT 预览生成失败，原文件仍保留在受保护存储中，可下载查看。管理员可检查文件是否损坏、格式是否过旧，必要时另存为 DOCX、PPTX 或 PDF 后重新上传。", tone="blue")

    add_heading(doc, "十、学生档案与归档资料", 1)
    add_body(doc, "学生档案用于集中保存学生基本信息和学术材料，不用于复杂的培养环节打卡。列表支持搜索、学位筛选、排序和分页，适合几十名学生长期使用。")
    add_heading(doc, "10.1 建档", 2)
    decimal_num += 1
    add_numbered(doc, "先在“成员管理”中创建或确认学生账号，并设置正确的学校身份。", decimal_num)
    add_numbered(doc, "管理员可一键生成学生档案；系统会检查该账号是否已有档案。", decimal_num)
    add_numbered(doc, "在档案中补充年级、研究题目、研究方向和导师信息。", decimal_num)
    add_numbered(doc, "有照片时优先显示照片；没有照片时自动使用姓名首字作为头像。", decimal_num)
    add_heading(doc, "10.2 上传归档资料", 2)
    add_bullet(doc, "支持开题报告、开题 PPT、中期报告、中期 PPT、毕业论文、答辩 PPT、发表论文、毕业交接材料和其他资料。", bullet_num)
    add_bullet(doc, "可一次选择或拖入多个文件，上传前可逐项修改资料名称和说明。", bullet_num)
    add_bullet(doc, "上传后仍可编辑标题、说明等信息；组内成员可查阅，删除权限更严格。", bullet_num)
    add_bullet(doc, "成员毕业后保留档案和历史资料，只需将成员状态改为“已毕业/离组”。", bullet_num)

    add_heading(doc, "十一、仪器平台使用说明", 1)
    add_body(doc, "仪器平台定位为组内设备信息目录，而非预约系统。它帮助成员快速了解设备是否可用、放在哪里、由谁负责以及如何使用。")
    decimal_num += 1
    add_numbered(doc, "新增仪器时填写名称，型号可选，补充状态、详细位置和负责人。", decimal_num)
    add_numbered(doc, "上传一张清晰设备图片，系统会自动压缩并在列表和详情页使用。", decimal_num)
    add_numbered(doc, "在使用说明中写明开关机、样品要求、注意事项和联系人等必要信息。", decimal_num)
    add_numbered(doc, "设备维护时将状态改为“维护中”；长期不用时改为“停用”。", decimal_num)
    add_numbered(doc, "删除属于低频操作，放在设备详情页中，操作前应再次确认。", decimal_num)
    add_callout(doc, "不在系统中记录的内容", "仪器预约、培训资格、复杂设备编号和实际使用台账不在本系统内维护，继续按课题组线下制度执行。", tone="gray")

    add_heading(doc, "十二、批量导入", 1)
    add_body(doc, "当账号、成员、资料、仪器或科研成果数量较多时，可使用 Excel 模板批量导入。建议先用 2—3 条数据试导入，确认格式无误后再导入完整数据。")
    add_table(
        doc,
        ["可批量导入的内容", "模板特点", "主要去重依据"],
        [
            ["账号", "账号名、姓名、身份、状态和系统权限", "账号名或邮箱"],
            ["公开团队成员", "姓名、身份头衔、研究方向、简介、排序，可带头像", "姓名"],
            ["内部资料", "资料信息表和附件压缩包", "标题及文件信息"],
            ["仪器", "说明页 + 数据页，可在 Excel 中嵌入图片", "仪器名称等信息"],
            ["论文", "以 GB/T 7714—2025 格式引文为主要输入", "DOI；无 DOI 时按题目和年份"],
            ["项目", "项目名称、编号、来源、周期和说明", "项目编号；无编号时按名称"],
            ["专利", "名称、专利号、发明人、日期和状态", "专利号；无专利号时按名称"],
            ["获奖", "奖项、日期、参与人员、说明，可带图片", "奖项名称和日期"],
        ],
        [2050, 4700, 2610],
        compact=True,
    )
    decimal_num += 1
    add_numbered(doc, "从对应页面下载最新模板，不要使用旧版本模板。", decimal_num)
    add_numbered(doc, "先阅读模板第一个工作表中的填写说明。", decimal_num)
    add_numbered(doc, "在数据工作表逐行填写，不要修改列名和工作表名称。", decimal_num)
    add_numbered(doc, "上传后等待结果提示，过程中不要重复点击上传。", decimal_num)
    add_numbered(doc, "核对新增、更新、跳过和错误行数量，再抽查几条前台结果。", decimal_num)

    add_heading(doc, "十三、文件和图片规则", 1)
    add_table(
        doc,
        ["文件类型", "当前建议上限", "使用建议"],
        [
            ["文档与归档资料", "200 MB", "适合开题报告、论文、PPT 和大型实验材料；慢速网络会显示进度"],
            ["普通图片", "20 MB", "正文插图、新闻封面、横幅和设备图片；上传后会自动压缩"],
            ["头像", "10 MB", "前端先裁剪固定比例，再压缩保存"],
            ["Logo", "5 MB", "建议使用清晰、背景干净的 PNG"],
            ["网站图标", "2 MB", "建议使用正方形 PNG 或 ICO"],
            ["Excel 批量导入", "50 MB", "不要在表格中嵌入过多超高清原图"],
        ],
        [1700, 1800, 5860],
        compact=True,
    )
    add_bullet(doc, "上传大文件时页面会显示真实进度；上传完成后还可能需要等待服务器生成预览。", bullet_num)
    add_bullet(doc, "文件名、文件大小和可执行操作会显示在前端，长文件名在手机上自动截断，不影响原文件。", bullet_num)
    add_bullet(doc, "受保护资料不能通过猜测文件网址绕过登录权限。", bullet_num)

    add_heading(doc, "十四、推荐的日常管理方式", 1)
    add_table(
        doc,
        ["频率", "建议工作", "建议负责人"],
        [
            ["随时", "发布新闻、补充新论文/项目/专利/获奖，更新仪器状态和学生资料", "网站编辑、对应业务管理员"],
            ["每月", "检查首页重点内容、失效外链、长期处于维护状态的设备和待审核账号", "系统管理员"],
            ["每学期", "更新成员身份、毕业/离组状态、学生档案和导师关系", "系统管理员或课题组秘书"],
            ["每季度", "执行数据库和媒体文件备份，抽查能否恢复", "服务器维护人员"],
            ["每年", "清理过时草稿，复核权限名单、域名证书和服务器空间", "课题组负责人 + 维护人员"],
        ],
        [1200, 5760, 2400],
        compact=True,
    )
    add_callout(doc, "推荐分工", "至少保留 1 名系统管理员、1—2 名网站编辑；资料和仪器可分别指定负责人。不要让所有成员都拥有管理员权限。", tone="green")

    add_heading(doc, "十五、数据安全、备份和更新", 1)
    add_body(doc, "网站程序、数据库和上传文件是三类不同资产。仅把程序上传到 GitHub，并不等于已经备份数据库和文件。")
    add_table(
        doc,
        ["资产", "保存内容", "维护要求"],
        [
            ["程序代码", "网页和功能程序", "通过 GitHub 保存版本，更新时在服务器运行部署脚本"],
            ["数据库", "账号、文字内容、分类、权限和记录", "定期运行数据库备份脚本，并把备份复制到服务器之外"],
            ["媒体文件", "头像、横幅、图片、PDF、Word、PPT 等", "定期运行媒体备份脚本，保护公开和内部文件目录"],
            ["服务器配置", "域名、密码、数据库连接等", "保存在服务器环境配置中，不上传到 GitHub，不在群聊公开"],
        ],
        [1500, 3600, 4260],
        compact=True,
    )
    add_heading(doc, "15.1 日常更新", 2)
    add_body(doc, "开发人员将修改推送到 GitHub 后，服务器维护人员进入项目目录并运行项目自带的部署脚本。部署脚本负责拉取更新、构建前后端、执行数据库变更并重新启动服务。普通内容更新不需要重新部署，直接在门户内容或内部平台保存即可。")
    add_heading(doc, "15.2 出现异常时的处理顺序", 2)
    decimal_num += 1
    add_numbered(doc, "先确认是所有页面打不开，还是只有某个文件、账号或页面异常。", decimal_num)
    add_numbered(doc, "记录发生时间、访问地址、操作步骤和页面提示，尽量保留截图。", decimal_num)
    add_numbered(doc, "若出现 502，通常表示网页入口暂时无法连接后台服务，应由服务器维护人员检查服务状态。", decimal_num)
    add_numbered(doc, "不要通过反复重新部署、删除数据库或重置全部账号来尝试修复。", decimal_num)
    add_numbered(doc, "重大操作前先做数据库和媒体文件备份。", decimal_num)

    add_heading(doc, "十六、常见问题", 1)
    faqs = [
        ("为什么创建了账号，却没有出现在官网团队成员里？", "内部账号和公开团队成员用途不同。账号用于登录，公开成员需在门户内容中单独维护。"),
        ("成员毕业后是否要删除账号？", "一般不删除。将成员状态改为“已毕业/离组”，保留学生档案和历史资料；确有需要时再停用登录。"),
        ("为什么 Word 或 PPT 不能立即预览？", "这类文件需要服务器转换。大文件或复杂文档耗时更长；失败时原文件仍可下载。"),
        ("为什么上传完成后还在等待？", "上传进度达到 100% 只代表文件已传到服务器，预览转换、图片处理或批量写入仍可能继续。"),
        ("为什么首页没有显示某条成果或成员？", "检查排序是否为 0、内容是否公开、是否已发布，以及首页展示数量是否已达到上限。"),
        ("普通成员能删除他人资料吗？", "不能。普通成员主要操作自己上传的内容，管理员或资料管理员拥有更高管理权限。"),
        ("内容保存后为什么图片没有立即变化？", "浏览器可能仍在使用旧缓存，可等待片刻或强制刷新；系统已通过更新时间尽量避免品牌图片长期缓存。"),
        ("网站是否支持手机？", "支持。公开导航、内部菜单、表单、分页、图片和资料预览都进行了手机适配。"),
        ("是否需要懂编程才能维护？", "不需要。日常内容、账号、资料、仪器和学生档案均在网页中维护；只有服务器部署和故障排查需要技术人员。"),
    ]
    for question, answer in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(question)
        set_run_font(r, size=11, bold=True, color=DEEP_GREEN)
        add_body(doc, answer, color=TEXT, after=5)

    add_heading(doc, "十七、系统检查清单", 1)
    add_checklist(
        doc,
        [
            ("正式访问地址可打开，电脑和手机均能正常访问", "通过"),
            ("管理员可用账号名和邮箱登录，并能正常退出", "通过"),
            ("首页横幅、简介、Logo、图标、地址和外链可在后台修改", "通过"),
            ("研究方向、团队成员、新闻和各类成果可新增、编辑、删除和排序", "通过"),
            ("图文详情可插入图片，列表不显示格式代码", "通过"),
            ("新闻、成员、研究方向和成果详情可显示浏览次数及更新时间", "通过"),
            ("内部资料可上传、搜索、排序、分页、预览、下载和编辑", "通过"),
            ("仪器可新建、编辑、搜索、排序、分页、批量导入和删除", "通过"),
            ("学生档案支持多导师、批量归档资料和在线预览", "通过"),
            ("账号、成员、资料、仪器和成果模板可下载并批量导入", "通过"),
            ("权限测试覆盖未登录、普通成员、各管理员和系统管理员", "通过"),
            ("服务器已配置数据库和媒体备份方式，维护人员掌握更新流程", "需现场确认"),
        ],
    )

    add_heading(doc, "十八、网站运行资料", 1)
    add_table(
        doc,
        ["资料项", "建议保管人", "注意事项"],
        [
            ["正式域名和服务器信息", "课题组负责人/指定维护人", "不要在公开文档中记录服务器密码"],
            ["系统管理员账号", "至少两名可信负责人", "使用强密码，不多人共用同一账号"],
            ["GitHub 仓库", "项目负责人", "用于保存程序版本，不存放密码和业务文件"],
            ["数据库备份", "指定维护人", "至少保留近期多份，并保存一份异地副本"],
            ["媒体文件备份", "指定维护人", "与数据库备份保持相近时间点"],
            ["批量导入模板", "各业务管理员", "以网站页面下载的最新版本为准"],
            ["本说明书", "网站管理员", "人员交接时一并移交"],
        ],
        [2100, 2400, 4860],
        compact=True,
    )
    add_callout(doc, "交接建议", "每次更换管理员时，先创建新管理员账号并验证可用，再停用旧账号；同时移交服务器、GitHub、域名、备份位置和本说明书。", tone="gold")

    add_heading(doc, "附录：快速操作索引", 1)
    add_table(
        doc,
        ["我要做什么", "从哪里进入", "关键提醒"],
        [
            ["修改首页文字或横幅", "内部平台 → 门户内容", "保存后到官网首页复核"],
            ["发布新闻", "门户内容 → 新闻活动", "先草稿预览，再发布"],
            ["添加论文/项目/专利/获奖", "门户内容 → 对应成果栏目", "数量多时使用最新 Excel 模板"],
            ["新增公开团队成员", "门户内容 → 团队成员", "排序为 0 时不公开展示"],
            ["创建登录账号", "成员管理 → 新建账号", "设置身份、审核状态和必要权限"],
            ["生成学生档案", "成员管理 → 对应学生账号", "一个账号只能有一个档案"],
            ["上传组内资料", "内部资料 → 上传资料", "当前分类会自动带入"],
            ["批量上传学生材料", "学生档案 → 选择学生 → 批量上传", "上传前逐项核对名称"],
            ["新增或维护仪器", "仪器平台", "预约和使用登记继续线下完成"],
            ["修改自己的资料", "右上角账号菜单 → 个人信息", "可修改账号名、姓名、头像和个人简介"],
            ["更新服务器版本", "由维护人员运行部署脚本", "更新前先确认备份"],
        ],
        [2800, 3300, 3260],
        compact=True,
    )

    doc.core_properties.title = "中农雨磷课题组网站使用与管理说明"
    doc.core_properties.subject = "网站功能、日常操作与运行管理说明"
    doc.core_properties.author = "中农雨磷课题组"
    doc.core_properties.keywords = "中农雨磷, LabHub, 网站说明, 内部平台, 网站管理"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
